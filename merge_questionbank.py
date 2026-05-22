# -*- coding: utf-8 -*-
"""
三文件整合脚本：
1. 读取"积分系统题库（2022）2024.7.8.docx"作为格式模板 + 题目源
2. 读取"综合气象观测.docx" 
3. 读取"县级综合气象业务考试系统题库-公共气象服务（总表）.xls"
去重、重编号、归类到"应知应会业务库"，输出"气象练习题库.docx"
"""
import re
import json
import copy
from collections import Counter, defaultdict, OrderedDict
from docx import Document
from docx.shared import Pt, Cm, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION_START
from docx.oxml.ns import qn
import xlrd

BASE_DIR = 'D:/W/AIproject/dataruanjina/html/'

# ==================== 解析文件1：积分系统题库 ====================
def parse_file1():
    doc = Document(BASE_DIR + '积分系统题库（2022）2024.7.8.docx')
    paras = [p.text.strip() for p in doc.paragraphs]

    questions = []
    chapter_names = ['公共气象服务', '气象预警预报', '综合气象观测', '综合气象保障']

    sections = []
    for i, txt in enumerate(paras):
        if txt == '应知应会业务库' or txt == '竞赛业务库':
            sections.append((txt, i))

    for si, (bank, bank_pos) in enumerate(sections):
        next_bank = sections[si + 1][1] if si + 1 < len(sections) else len(paras)
        ch_positions = []
        for i in range(bank_pos, next_bank):
            if paras[i] in chapter_names:
                ch_positions.append((paras[i], i))

        for ci, (ch, ch_pos) in enumerate(ch_positions):
            next_ch = ch_positions[ci + 1][1] if ci + 1 < len(ch_positions) else next_bank
            type_positions = []
            for i in range(ch_pos, next_ch):
                if re.match(r'^[一二三][\、．\.]', paras[i]):
                    type_positions.append(i)

            for ti, type_pos in enumerate(type_positions):
                next_type = type_positions[ti + 1] if ti + 1 < len(type_positions) else next_ch
                type_text = paras[type_pos]
                if '单选' in type_text:
                    q_type = '单选题'
                elif '多选' in type_text:
                    q_type = '多选题'
                elif '判断' in type_text:
                    q_type = '判断题'
                else:
                    continue

                # 收集题目
                i = type_pos + 1
                while i < next_type:
                    txt = paras[i]
                    m = re.match(r'^\d+[\.\．](.*)', txt, re.DOTALL)
                    if m:
                        content = m.group(1).strip()
                        content, answer = extract_answer(content)

                        options = []
                        j = i + 1
                        while j < next_type:
                            opt_txt = paras[j]
                            if re.match(r'^[A-H][\、\.\．\s]', opt_txt):
                                opt_match = re.match(r'^([A-H])[\、\.\．\s](.*)', opt_txt)
                                if opt_match:
                                    options.append({'label': opt_match.group(1), 'content': opt_match.group(2).strip()})
                                j += 1
                            elif re.match(r'^\d+[\.\．]', opt_txt):
                                break
                            elif opt_txt in ('对', '错'):
                                break
                            else:
                                split_match = re.match(r'^(.*?)([A-H][\、\.\．\s].*)$', opt_txt)
                                if split_match and options:
                                    options[-1]['content'] += split_match.group(1).strip()
                                    opt_txt = split_match.group(2)
                                    continue
                                if options:
                                    options[-1]['content'] += opt_txt
                                j += 1

                        if q_type == '判断题':
                            if j < next_type and paras[j] in ('对', '错'):
                                answer = paras[j]
                                j += 1

                        questions.append({
                            'source': 'file1',
                            'chapter': ch,
                            'type': q_type,
                            'content': content,
                            'answer': answer or '',
                            'options': options,
                            'analysis': ''
                        })
                        i = j
                    else:
                        i += 1

    return questions


def extract_answer(content):
    patterns = [
        (r'_{2,}([A-H]+)\s*_{2,}', True),
        (r'_{2,}([A-H]+)_{2,}', True),
        (r'_([A-H]+)_', True),
        (r'\(\s*([A-H]+)\s*\)', False),
        (r'（\s*([A-H]+)\s*）', False),
        (r'【\s*([A-H]+)\s*】', False),
        (r'\(\s*([A-Za-z]{2,})\s*\)', False),
        (r'（\s*([A-Za-z]{2,})\s*）', False),
        (r'_{2,}([A-Za-z]{2,})_{2,}', True),
        (r'_([A-Za-z]{2,})_', True),
        (r'【\s*([A-Za-z]{2,})\s*】', False),
        (r'\s{2,}([A-H])\s{2,}', False),
        (r'\s([A-H])\s_', False),
    ]
    for pat, use_underscore in patterns:
        m = re.search(pat, content)
        if m:
            answer = m.group(1).upper()
            if use_underscore:
                content = re.sub(pat, '____', content, count=1)
            else:
                content = re.sub(pat, '（）', content, count=1)
            return content, answer
    return content, None


# ==================== 解析文件2：综合气象观测.docx ====================
def parse_file2():
    doc = Document(BASE_DIR + '综合气象观测.docx')
    paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    questions = []
    current_type = None

    i = 0
    while i < len(paras):
        txt = paras[i]

        if txt in ('单选题', '多选题', '判断题'):
            current_type = txt
            i += 1
            continue

        if current_type is None:
            i += 1
            continue

        # 判断题格式: "对1.题干" 或 "错1.题干"
        m_judge = re.match(r'^(对|错)(\d+)[\.\．]?(.*)', txt, re.DOTALL)
        # 选择题格式: "A1.题干" 答案是单个大写字母
        m_choice = re.match(r'^([A-H])(\d+)[\.\．]?(.*)', txt, re.DOTALL)
        # 多选题格式: "C/D171.题干" 或 "AB12.题干"
        m_multi = re.match(r'^([A-H][/A-H]+)(\d+)[\.\．]?(.*)', txt, re.DOTALL)

        if current_type == '判断题' and m_judge:
            answer = m_judge.group(1)
            content = m_judge.group(3).strip()
            questions.append({
                'source': 'file2', 'chapter': '综合气象观测', 'type': '判断题',
                'content': content, 'answer': answer, 'options': [], 'analysis': ''
            })
            i += 1
        elif m_multi and not m_choice:
            # 多选答案如 C/D
            ans_raw = m_multi.group(1).replace('/', '')
            content = m_multi.group(3).strip()
            options = []
            # 处理同行选项
            content, options = parse_inline_options(content)
            if not options:
                # 从后续行收集选项
                j = i + 1
                while j < len(paras):
                    opt_txt = paras[j]
                    if re.match(r'^([A-H])[、\.\．\s]', opt_txt):
                        om = re.match(r'^([A-H])[、\.\．\s](.*)', opt_txt)
                        if om:
                            options.append({'label': om.group(1), 'content': om.group(2).strip()})
                        j += 1
                    elif re.match(r'^([A-H]+|对|错)\d+[\.\．]?', opt_txt) or opt_txt in ('单选题','多选题','判断题'):
                        break
                    else:
                        if options:
                            options[-1]['content'] += opt_txt
                        j += 1
                i = j
            else:
                i += 1
            questions.append({
                'source': 'file2', 'chapter': '综合气象观测', 'type': current_type,
                'content': content, 'answer': ans_raw, 'options': options, 'analysis': ''
            })
        elif m_choice:
            answer = m_choice.group(1)
            content = m_choice.group(3).strip()
            options = []
            content, options = parse_inline_options(content)
            if not options:
                j = i + 1
                while j < len(paras):
                    opt_txt = paras[j]
                    if re.match(r'^([A-H])[、\.\．\s]', opt_txt):
                        om = re.match(r'^([A-H])[、\.\．\s](.*)', opt_txt)
                        if om:
                            options.append({'label': om.group(1), 'content': om.group(2).strip()})
                        j += 1
                    elif re.match(r'^([A-H]+|对|错)\d+[\.\．]?', opt_txt) or opt_txt in ('单选题','多选题','判断题'):
                        break
                    else:
                        if options:
                            options[-1]['content'] += opt_txt
                        j += 1
                i = j
            else:
                i += 1
            questions.append({
                'source': 'file2', 'chapter': '综合气象观测', 'type': current_type,
                'content': content, 'answer': answer, 'options': options, 'analysis': ''
            })
        else:
            i += 1

    return questions


def parse_inline_options(content):
    """从题干中提取同行选项，返回 (cleaned_content, options)"""
    options = []
    # 查找 "A、xxx B、xxx" 格式
    parts = re.split(r'\s+([A-F])[、\.\．\s]', content)
    if len(parts) >= 3:
        content = parts[0].strip()
        idx = 1
        while idx + 1 < len(parts):
            label = parts[idx]
            opt_content = parts[idx + 1].strip().rstrip('；;')
            if label in 'ABCDEF' and opt_content:
                options.append({'label': label, 'content': opt_content})
            idx += 2
    return content, options


# ==================== 解析文件3：公共气象服务.xls ====================
def parse_file3():
    wb = xlrd.open_workbook(BASE_DIR + '县级综合气象业务考试系统题库-公共气象服务（总表）.xls')
    sheet = wb.sheet_by_index(0)

    questions = []
    for r in range(2, sheet.nrows):
        chapter_raw = str(sheet.cell_value(r, 0)).strip()
        q_type = str(sheet.cell_value(r, 1)).strip()
        content = str(sheet.cell_value(r, 2)).strip()
        answer = str(sheet.cell_value(r, 15)).strip() if sheet.ncols > 15 else ''
        analysis = str(sheet.cell_value(r, 16)).strip() if sheet.ncols > 16 else ''

        if not content or not q_type:
            continue

        # 提取章节名
        if '/' in chapter_raw:
            chapter = chapter_raw.split('/')[-1].strip()
        else:
            chapter = chapter_raw or '公共气象服务'

        # 标准化题型
        if '单选' in q_type:
            q_type = '单选题'
        elif '多选' in q_type:
            q_type = '多选题'
        elif '判断' in q_type:
            q_type = '判断题'
        else:
            continue

        # 标准化答案
        if q_type == '判断题':
            if answer in ('对', '错'):
                pass
            else:
                answer = '对' if '对' in answer else ('错' if '错' in answer else '')
        else:
            # 多选题答案可能有逗号分隔
            answer = answer.replace(',', '').replace('，', '').replace('/', '')
            answer = ''.join(sorted(c for c in answer.upper() if c in 'ABCDEFGHIJ'))

        # 收集选项
        options = []
        opt_labels = ['A', 'B', 'C', 'D', 'E', 'F']
        for idx, label in enumerate(opt_labels):
            col = 4 + idx * 2  # A=col4, B=col6, C=col8, D=col10, E=col12, F=col14
            if col < sheet.ncols:
                opt_content = str(sheet.cell_value(r, col)).strip()
                if opt_content:
                    options.append({'label': label, 'content': opt_content})

        if not analysis:
            analysis = ''

        questions.append({
            'source': 'file3',
            'chapter': '公共气象服务',
            'type': q_type,
            'content': content,
            'answer': answer,
            'options': options,
            'analysis': analysis
        })

    return questions


# ==================== 去重与整合 ====================
def deduplicate(all_qs):
    seen = set()
    unique = []
    dup_count = 0
    for q in all_qs:
        key = (q['type'], q['content'])
        if key not in seen:
            seen.add(key)
            unique.append(q)
        else:
            dup_count += 1
    return unique, dup_count


# ==================== 生成Word文档 ====================
def generate_docx(questions, template_path, output_path):
    template = Document(template_path)
    doc = Document()

    # 复制模板页面设置
    for section in template.sections:
        new_section = doc.add_section()
        new_section.page_width = section.page_width
        new_section.page_height = section.page_height
        new_section.left_margin = section.left_margin
        new_section.right_margin = section.right_margin
        new_section.top_margin = section.top_margin
        new_section.bottom_margin = section.bottom_margin

    # 删除默认段落（去除首空白页）
    while doc.paragraphs:
        p = doc.paragraphs[0]
        p._element.getparent().remove(p._element)

    # 字体设置辅助函数
    def set_font(run, font_name='楷体', size=Pt(10.5), bold=False):
        run.font.name = font_name
        run.font.size = size
        run.bold = bold
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = rPr.makeelement(qn('w:rFonts'), {})
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:eastAsia'), font_name)

    def set_paragraph_format(paragraph, line_spacing=1.0, space_before=Pt(0), space_after=Pt(0)):
        pf = paragraph.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = line_spacing
        pf.space_before = space_before
        pf.space_after = space_after

    def add_bookmark(paragraph, bookmark_name):
        pPr = paragraph._element.get_or_add_pPr()
        bmStart = pPr.makeelement(qn('w:bookmarkStart'), {
            qn('w:id'): str(hash(bookmark_name) % 10000),
            qn('w:name'): bookmark_name
        })
        pPr.append(bmStart)
        bmEnd = paragraph._element.makeelement(qn('w:bookmarkEnd'), {
            qn('w:id'): str(hash(bookmark_name) % 10000)
        })
        paragraph._element.append(bmEnd)

    def add_hyperlink(paragraph, bookmark_name, text, font_name='楷体', size=Pt(12)):
        hyperlink = paragraph._element.makeelement(qn('w:hyperlink'), {
            qn('w:anchor'): bookmark_name
        })
        run_elem = hyperlink.makeelement(qn('w:r'), {})
        rPr = run_elem.makeelement(qn('w:rPr'), {})
        rFonts = rPr.makeelement(qn('w:rFonts'), {
            qn('w:ascii'): font_name,
            qn('w:hAnsi'): font_name,
            qn('w:eastAsia'): font_name
        })
        rPr.append(rFonts)
        sz = rPr.makeelement(qn('w:sz'), {qn('w:val'): str(int(size.pt * 2))})
        szCs = rPr.makeelement(qn('w:szCs'), {qn('w:val'): str(int(size.pt * 2))})
        rPr.append(sz)
        rPr.append(szCs)
        color = rPr.makeelement(qn('w:color'), {qn('w:val'): '0563C1'})
        rPr.append(color)
        u = rPr.makeelement(qn('w:u'), {qn('w:val'): 'single'})
        rPr.append(u)
        run_elem.append(rPr)
        t_elem = run_elem.makeelement(qn('w:t'), {})
        t_elem.text = text
        run_elem.append(t_elem)
        hyperlink.append(run_elem)
        paragraph._element.append(hyperlink)

    # 按章节和题型分组
    chapter_order = ['公共气象服务', '气象预警预报', '综合气象观测', '综合气象保障']
    type_order = ['单选题', '多选题', '判断题']

    grouped = OrderedDict()
    for ch in chapter_order:
        ch_qs = [q for q in questions if q['chapter'] == ch]
        if not ch_qs:
            continue
        grouped[ch] = OrderedDict()
        for t in type_order:
            t_qs = [q for q in ch_qs if q['type'] == t]
            if t_qs:
                grouped[ch][t] = t_qs

    type_label_map = {'单选题': '一、单选题', '多选题': '二、多选题', '判断题': '三、判断题'}

    # ========== 第一部分：目录页（单栏） ==========
    # 目录页标题
    p = doc.add_paragraph()
    run = p.add_run('目  录')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(run, '楷体', Pt(18), bold=True)
    set_paragraph_format(p, space_after=Pt(12))

    for ch, types in grouped.items():
        # 章节条目（一级）
        p_ch = doc.add_paragraph()
        set_paragraph_format(p_ch, space_before=Pt(6), space_after=Pt(2))
        bm_ch = f'toc_{ch}'
        add_hyperlink(p_ch, bm_ch, ch, '楷体', Pt(14))

        for t in types:
            # 题型条目（二级，缩进）
            p_t = doc.add_paragraph()
            p_t.paragraph_format.left_indent = Cm(1)
            set_paragraph_format(p_t, space_before=Pt(2), space_after=Pt(1))
            bm_t = f'toc_{ch}_{t}'
            add_hyperlink(p_t, bm_t, type_label_map.get(t, t), '楷体', Pt(12))

    # 目录后分节符（下一页），正文两栏
    new_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    for section in template.sections:
        new_section.page_width = section.page_width
        new_section.page_height = section.page_height
        new_section.left_margin = section.left_margin
        new_section.right_margin = section.right_margin
        new_section.top_margin = section.top_margin
        new_section.bottom_margin = section.bottom_margin
        break

    # ========== 第二部分：正文（两栏） ==========
    # 设置两栏排版（仅正文section）
    sectPr = new_section._sectPr
    cols = sectPr.find(qn('w:cols'))
    if cols is None:
        cols = sectPr.makeelement(qn('w:cols'), {})
        sectPr.append(cols)
    cols.set(qn('w:num'), '2')

    # 标题：题库名
    p = doc.add_paragraph()
    run = p.add_run('应知应会业务库')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(run, '楷体', Pt(16), bold=True)
    set_paragraph_format(p)

    for ch, types in grouped.items():
        # 章节标题 + 书签
        p = doc.add_paragraph()
        run = p.add_run(ch)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(run, '楷体', Pt(14), bold=True)
        set_paragraph_format(p)
        add_bookmark(p, f'toc_{ch}')

        for t, qs in types.items():
            # 题型标题 + 书签
            p = doc.add_paragraph()
            run = p.add_run(type_label_map.get(t, t))
            set_font(run, '楷体', Pt(12), bold=True)
            set_paragraph_format(p)
            add_bookmark(p, f'toc_{ch}_{t}')

            num = 1
            for q in qs:
                if t == '判断题':
                    p = doc.add_paragraph()
                    run = p.add_run(f'{num}.{q["content"]}')
                    set_font(run, '楷体', Pt(10.5))
                    set_paragraph_format(p)
                    p2 = doc.add_paragraph()
                    run2 = p2.add_run(q['answer'])
                    set_font(run2, '楷体', Pt(10.5))
                    set_paragraph_format(p2)
                else:
                    content_with_ans = q['content']
                    ans_str = q['answer']
                    if '____' in content_with_ans:
                        content_with_ans = content_with_ans.replace('____', f'__{ans_str}__', 1)
                    elif '（）' in content_with_ans:
                        content_with_ans = content_with_ans.replace('（）', f'（{ans_str}）', 1)
                    else:
                        content_with_ans = f'{content_with_ans}（{ans_str}）'

                    p = doc.add_paragraph()
                    run = p.add_run(f'{num}.{content_with_ans}')
                    set_font(run, '楷体', Pt(10.5))
                    set_paragraph_format(p)

                    for opt in q['options']:
                        p_opt = doc.add_paragraph()
                        run_opt = p_opt.add_run(f'{opt["label"]}、{opt["content"]}')
                        set_font(run_opt, '楷体', Pt(10.5))
                        set_paragraph_format(p_opt)

                if q.get('analysis') and q['analysis'] != '无' and q['analysis'].strip():
                    pa = doc.add_paragraph()
                    run_a = pa.add_run(f'解析：{q["analysis"]}')
                    set_font(run_a, '楷体', Pt(9))
                    set_paragraph_format(pa)

                num += 1

    # ========== 添加页脚页码（1/N格式）到所有section ==========
    for i, section in enumerate(doc.sections):
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 当前页码域
        run1 = fp.add_run()
        fldChar1 = run1._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'})
        run1._element.append(fldChar1)

        run2 = fp.add_run()
        instrText2 = run2._element.makeelement(qn('w:instrText'), {})
        instrText2.text = ' PAGE '
        run2._element.append(instrText2)

        run3 = fp.add_run()
        fldChar3 = run3._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'})
        run3._element.append(fldChar3)

        # 分隔符 /
        fp.add_run('/')

        # 总页码域
        run4 = fp.add_run()
        fldChar4 = run4._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'})
        run4._element.append(fldChar4)

        run5 = fp.add_run()
        instrText5 = run5._element.makeelement(qn('w:instrText'), {})
        instrText5.text = ' NUMPAGES '
        run5._element.append(instrText5)

        run6 = fp.add_run()
        fldChar6 = run6._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'})
        run6._element.append(fldChar6)

    doc.save(output_path)


# ==================== 主流程 ====================
def main():
    print('=== 解析文件1: 积分系统题库 ===')
    qs1 = parse_file1()
    print(f'  解析出 {len(qs1)} 题')

    print('=== 解析文件2: 综合气象观测 ===')
    qs2 = parse_file2()
    print(f'  解析出 {len(qs2)} 题')

    print('=== 解析文件3: 公共气象服务.xls ===')
    qs3 = parse_file3()
    print(f'  解析出 {len(qs3)} 题')

    all_qs = qs1 + qs2 + qs3
    print(f'\n=== 合计: {len(all_qs)} 题 ===')

    # 去重
    unique_qs, dup_count = deduplicate(all_qs)
    print(f'去重: 移除 {dup_count} 道重复题, 保留 {len(unique_qs)} 题')

    # 修复：单选/多选题答案不是A-H字母组合 → 填空题
    # 修复：判断题答案不在(对,错) → 空答案尝试推断或标记
    fill_count = 0
    for q in unique_qs:
        a = q['answer'].strip()
        if q['type'] in ('单选题', '多选题'):
            if not a:
                q['type'] = '填空题'
                fill_count += 1
            elif not all(c in 'ABCDEFGHIJ' for c in a.upper()):
                q['type'] = '填空题'
                fill_count += 1
        elif q['type'] == '判断题':
            if a not in ('对', '错'):
                if '对' in a:
                    q['answer'] = '对'
                elif '错' in a:
                    q['answer'] = '错'
                else:
                    q['type'] = '填空题'
                    fill_count += 1
    if fill_count:
        print(f'重新归类为填空题: {fill_count}题')

    # 补充解析
    for q in unique_qs:
        if not q.get('analysis') or q['analysis'].strip() == '':
            q['analysis'] = '无'

    # 统计
    print('\n=== 最终题目统计 ===')
    for ch in ['公共气象服务', '气象预警预报', '综合气象观测', '综合气象保障']:
        ch_qs = [q for q in unique_qs if q['chapter'] == ch]
        if not ch_qs:
            continue
        for t in ['单选题', '多选题', '判断题', '填空题']:
            t_qs = [q for q in ch_qs if q['type'] == t]
            if t_qs:
                print(f'  {ch}/{t}: {len(t_qs)}题')
        print(f'  {ch} 小计: {len(ch_qs)}题')
    print(f'  总计: {len(unique_qs)}题')

    # 生成Word
    print('\n=== 生成Word文档 ===')
    generate_docx(
        unique_qs,
        BASE_DIR + '积分系统题库（2022）2024.7.8.docx',
        BASE_DIR + '气象练习题库.docx'
    )
    print('已保存: 气象练习题库.docx')

    # 保存JSON供前端使用
    # 重新编号并输出JSON
    q_id = 0
    json_qs = []
    for ch in ['公共气象服务', '气象预警预报', '综合气象观测', '综合气象保障']:
        ch_qs = [q for q in unique_qs if q['chapter'] == ch]
        for t in ['单选题', '多选题', '判断题', '填空题']:
            t_qs = [q for q in ch_qs if q['type'] == t]
            for q in t_qs:
                json_qs.append({
                    'id': q_id,
                    'bank': '应知应会业务库',
                    'chapter': ch,
                    'type': q['type'],
                    'content': q['content'],
                    'answer': q['answer'],
                    'options': q['options'],
                    'analysis': q.get('analysis', '')
                })
                q_id += 1

    with open(BASE_DIR + 'questions.json', 'w', encoding='utf-8') as f:
        json.dump(json_qs, f, ensure_ascii=False, indent=2)
    print(f'已更新: questions.json ({len(json_qs)}题)')


if __name__ == '__main__':
    main()
